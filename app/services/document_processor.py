"""
Document processing service for extracting text from various file formats.
Supports PDF, DOCX, TXT, and Markdown files with metadata extraction.
"""

import os
import uuid
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import PyPDF2
import docx

from app.core.config import settings
from app.utils.text_splitter import RecursiveCharacterTextSplitter
from app.core.logging import get_logger
from app.core.exceptions import FileProcessingError, ValidationError
from app.models.schemas import Document, DocumentChunk, DocumentMetadata, DocumentType, ProcessingStatus

logger = get_logger(__name__)


class DocumentProcessor:
    """Service for processing and extracting text from documents."""
    
    def __init__(self):
        """Initialize the document processor."""
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.max_chunk_size,
            chunk_overlap=settings.chunk_overlap,
            separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""]
        )
        
        # Ensure upload directory exists
        os.makedirs(settings.uploads_dir, exist_ok=True)
    
    async def process_document(
        self,
        file_path: str,
        filename: str,
        file_size: int
    ) -> Document:
        """
        Process a document and extract text and metadata.
        
        Args:
            file_path: Path to the uploaded file
            filename: Original filename
            file_size: File size in bytes
            
        Returns:
            Processed document with metadata
            
        Raises:
            FileProcessingError: If document processing fails
            ValidationError: If file format is not supported
        """
        try:
            logger.info(
                "Starting document processing",
                filename=filename,
                file_size=file_size,
                file_path=file_path
            )
            
            # Generate document ID
            document_id = str(uuid.uuid4())
            
            # Validate file type
            file_extension = Path(filename).suffix.lower()
            if file_extension not in settings.supported_file_types:
                raise ValidationError(
                    f"Unsupported file type: {file_extension}",
                    details={"supported_types": settings.supported_file_types}
                )
            
            # Extract file type
            doc_type = self._get_document_type(file_extension)
            
            # Extract text and metadata
            text_content, metadata = await self._extract_text_and_metadata(
                file_path, filename, file_size, doc_type
            )
            
            # Create document
            document = Document(
                id=document_id,
                filename=filename,
                file_path=file_path,
                metadata=metadata,
                processing_status=ProcessingStatus.COMPLETED
            )
            
            logger.info(
                "Document processing completed",
                document_id=document_id,
                filename=filename,
                text_length=len(text_content)
            )
            
            return document
            
        except (ValidationError, FileProcessingError):
            raise
        except Exception as e:
            logger.error(
                "Unexpected error during document processing",
                filename=filename,
                error=str(e),
                exc_info=True
            )
            raise FileProcessingError(
                f"Failed to process document: {str(e)}",
                details={"filename": filename, "error": str(e)}
            )
    
    async def create_chunks(
        self,
        document: Document,
        text_content: str
    ) -> List[DocumentChunk]:
        """
        Split document text into chunks.
        
        Args:
            document: Document instance
            text_content: Extracted text content
            
        Returns:
            List of document chunks
        """
        try:
            logger.info(
                "Creating document chunks",
                document_id=document.id,
                text_length=len(text_content)
            )
            
            # Split text into chunks
            text_chunks = self.text_splitter.split_text(text_content)
            
            # Create chunk objects
            chunks = []
            for i, chunk_text in enumerate(text_chunks):
                chunk = DocumentChunk(
                    chunk_id=f"{document.id}_chunk_{i}",
                    document_id=document.id,
                    content=chunk_text.strip(),
                    chunk_index=i,
                    metadata={
                        "filename": document.filename,
                        "file_type": document.metadata.file_type,
                        "chunk_length": len(chunk_text)
                    }
                )
                chunks.append(chunk)
            
            logger.info(
                "Document chunks created",
                document_id=document.id,
                chunk_count=len(chunks)
            )
            
            return chunks
            
        except Exception as e:
            logger.error(
                "Error creating document chunks",
                document_id=document.id,
                error=str(e),
                exc_info=True
            )
            raise FileProcessingError(
                f"Failed to create document chunks: {str(e)}",
                details={"document_id": document.id, "error": str(e)}
            )
    
    def _get_document_type(self, file_extension: str) -> DocumentType:
        """Get document type from file extension."""
        type_mapping = {
            ".pdf": DocumentType.PDF,
            ".txt": DocumentType.TEXT,
            ".docx": DocumentType.DOCX,
            ".md": DocumentType.MARKDOWN
        }
        return type_mapping.get(file_extension, DocumentType.TEXT)
    
    async def _extract_text_and_metadata(
        self,
        file_path: str,
        filename: str,
        file_size: int,
        doc_type: DocumentType
    ) -> Tuple[str, DocumentMetadata]:
        """
        Extract text content and metadata from a file.
        
        Args:
            file_path: Path to the file
            filename: Original filename
            file_size: File size in bytes
            doc_type: Document type
            
        Returns:
            Tuple of (text_content, metadata)
        """
        try:
            if doc_type == DocumentType.PDF:
                return await self._extract_pdf(file_path, filename, file_size)
            elif doc_type == DocumentType.DOCX:
                return await self._extract_docx(file_path, filename, file_size)
            elif doc_type in [DocumentType.TEXT, DocumentType.MARKDOWN]:
                return await self._extract_text(file_path, filename, file_size)
            else:
                raise ValidationError(f"Unsupported document type: {doc_type}")
                
        except Exception as e:
            logger.error(
                "Error extracting text and metadata",
                filename=filename,
                doc_type=doc_type,
                error=str(e),
                exc_info=True
            )
            raise FileProcessingError(
                f"Failed to extract text from {filename}: {str(e)}"
            )
    
    async def _extract_pdf(
        self,
        file_path: str,
        filename: str,
        file_size: int
    ) -> Tuple[str, DocumentMetadata]:
        """Extract text and metadata from PDF file."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract text
                text_content = ""
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
                
                # Extract metadata
                pdf_info = pdf_reader.metadata or {}
                
                metadata = DocumentMetadata(
                    filename=filename,
                    file_size=file_size,
                    file_type=DocumentType.PDF,
                    mime_type="application/pdf",
                    page_count=len(pdf_reader.pages),
                    title=pdf_info.get("/Title"),
                    author=pdf_info.get("/Author"),
                    subject=pdf_info.get("/Subject")
                )
                
                return text_content.strip(), metadata
                
        except Exception as e:
            raise FileProcessingError(f"Failed to process PDF file: {str(e)}")
    
    async def _extract_docx(
        self,
        file_path: str,
        filename: str,
        file_size: int
    ) -> Tuple[str, DocumentMetadata]:
        """Extract text and metadata from DOCX file."""
        try:
            doc = docx.Document(file_path)
            
            # Extract text
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract metadata
            core_props = doc.core_properties
            
            metadata = DocumentMetadata(
                filename=filename,
                file_size=file_size,
                file_type=DocumentType.DOCX,
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                title=core_props.title,
                author=core_props.author,
                subject=core_props.subject,
                keywords=core_props.keywords.split(",") if core_props.keywords else None
            )
            
            return text_content.strip(), metadata
            
        except Exception as e:
            raise FileProcessingError(f"Failed to process DOCX file: {str(e)}")
    
    async def _extract_text(
        self,
        file_path: str,
        filename: str,
        file_size: int
    ) -> Tuple[str, DocumentMetadata]:
        """Extract text from plain text or markdown file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            text_content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text_content = file.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if text_content is None:
                raise FileProcessingError(f"Could not decode text file: {filename}")
            
            # Determine file type
            file_extension = Path(filename).suffix.lower()
            doc_type = DocumentType.MARKDOWN if file_extension == ".md" else DocumentType.TEXT
            mime_type = "text/markdown" if doc_type == DocumentType.MARKDOWN else "text/plain"
            
            metadata = DocumentMetadata(
                filename=filename,
                file_size=file_size,
                file_type=doc_type,
                mime_type=mime_type
            )
            
            return text_content.strip(), metadata
            
        except Exception as e:
            raise FileProcessingError(f"Failed to process text file: {str(e)}")
    
    async def validate_file(self, file_path: str, filename: str) -> None:
        """
        Validate uploaded file.
        
        Args:
            file_path: Path to the file
            filename: Original filename
            
        Raises:
            ValidationError: If file validation fails
        """
        try:
            # Check if file exists
            if not os.path.exists(file_path):
                raise ValidationError(f"File not found: {filename}")
            
            # Check file size
            file_size = os.path.getsize(file_path)
            if file_size > settings.max_file_size:
                raise ValidationError(
                    f"File size exceeds maximum allowed size of {settings.max_file_size} bytes",
                    details={
                        "file_size": file_size,
                        "max_size": settings.max_file_size
                    }
                )
            
            # Check file extension
            file_extension = Path(filename).suffix.lower()
            if file_extension not in settings.supported_file_types:
                raise ValidationError(
                    f"Unsupported file type: {file_extension}",
                    details={"supported_types": settings.supported_file_types}
                )
            
            logger.info(
                "File validation successful",
                filename=filename,
                file_size=file_size,
                file_extension=file_extension
            )
            
        except ValidationError:
            raise
        except Exception as e:
            logger.error(
                "Error validating file",
                filename=filename,
                error=str(e),
                exc_info=True
            )
            raise ValidationError(f"File validation failed: {str(e)}")


# Global document processor instance
document_processor = DocumentProcessor()
